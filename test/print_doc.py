from docx import Document

def create_doc_file(path):
    doc = Document()
    
    doc.add_heading("Document Testing", level=1)
    doc.add_paragraph("This is an example paragraph")

    doc.save(path)

file_path = "test.docx"
create_doc_file(file_path)